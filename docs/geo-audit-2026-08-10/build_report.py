#!/usr/bin/env python3
"""Build HTML and DOCX from the single Markdown source (kami style)."""

from __future__ import annotations

import html as html_mod
import re
import sys
from pathlib import Path

import markdown

BASE = Path(__file__).resolve().parent
MD_PATH = BASE / "ketsuin-geo-page-audit.md"
HTML_PATH = BASE / "ketsuin-geo-page-audit.html"
DOCX_PATH = BASE / "ketsuin-geo-page-audit.docx"
CSS_PATH = Path("/Users/lizhuo/.agents/skills/yao-geo-page-audit/templates/report.css")


def build_html(md_text: str) -> None:
    md = markdown.Markdown(extensions=["tables", "fenced_code", "toc"])
    body = md.convert(md_text)

    def toc_items(tokens):
        out = []
        for tok in tokens:
            out.append((tok["level"], tok["id"], tok["name"]))
            out.extend(toc_items(tok.get("children", [])))
        return out

    links = []
    for level, anchor, name in toc_items(md.toc_tokens):
        if level <= 2:
            links.append(f'<li><a href="#{anchor}">{html_mod.escape(name)}</a></li>')
    toc_html = '<nav id="TOC"><ul>' + "".join(links) + "</ul></nav>"

    css = CSS_PATH.read_text(encoding="utf-8")
    extra = "html,body{-webkit-print-color-adjust:exact;print-color-adjust:exact}"
    page = (
        "<!DOCTYPE html>\n<html lang=\"zh-CN\"><head><meta charset=\"utf-8\">"
        "<meta name=\"viewport\" content=\"width=device-width,initial-scale=1\">"
        "<title>Ketsuin 結印 GEO 页面诊断报告</title>"
        f"<style>{css}{extra}</style></head><body>{toc_html}{body}</body></html>"
    )
    HTML_PATH.write_text(page, encoding="utf-8")


# ---------------- DOCX ----------------

def add_runs(paragraph, text: str) -> None:
    """Render inline **bold** and `code` and [text](url) into runs."""
    from docx.oxml.ns import qn

    pos = 0
    pattern = re.compile(r"(\*\*.+?\*\*|`.+?`|\[[^\]]+\]\([^)]*\))")
    for m in pattern.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos : m.start()])
        token = m.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Menlo"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Menlo")
        else:  # markdown link -> visible text only
            label = re.match(r"\[([^\]]+)\]", token).group(1)
            paragraph.add_run(label)
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def set_table_fixed(table, total_width_twips: int) -> None:
    from docx.oxml.ns import qn
    from docx.shared import Twips

    tbl = table._tbl
    tblPr = tbl.tblPr
    layout = tblPr.find(qn("w:tblLayout"))
    if layout is None:
        layout = tblPr.makeelement(qn("w:tblLayout"), {})
        tblPr.append(layout)
    layout.set(qn("w:type"), "fixed")
    ncols = len(table.columns)
    per = total_width_twips // ncols
    for col in table.columns:
        col.width = Twips(per)
    for row in table.rows:
        for cell in row.cells:
            cell.width = Twips(per)


def build_docx(md_text: str) -> None:
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE

    doc = Document()
    # code style so polish_docx.py can recognize it
    styles = doc.styles
    try:
        code_style = styles.add_style("SourceCode", WD_STYLE_TYPE.PARAGRAPH)
        code_style.font.name = "Menlo"
        code_style.font.size = None
    except ValueError:
        pass

    lines = md_text.splitlines()
    i = 0
    in_code = False
    code_buf: list[str] = []

    def flush_code():
        nonlocal code_buf
        for line in code_buf:
            p = doc.add_paragraph(style="SourceCode")
            p.add_run(line if line else " ")
        code_buf = []

    while i < len(lines):
        line = lines[i]
        if line.strip().startswith("```"):
            if in_code:
                flush_code()
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code_buf.append(line)
            i += 1
            continue

        stripped = line.strip()
        if not stripped:
            i += 1
            continue

        if stripped.startswith("|"):
            rows = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                rows.append(lines[i].strip())
                i += 1
            cells = [
                [c.strip() for c in row.strip("|").split("|")]
                for row in rows
                if not re.fullmatch(r"\|[\s\-:|]+\|", row)
            ]
            table = doc.add_table(rows=len(cells), cols=len(cells[0]))
            table.style = "Table Grid"
            for r, row_cells in enumerate(cells):
                for c, value in enumerate(row_cells):
                    cell_p = table.rows[r].cells[c].paragraphs[0]
                    add_runs(cell_p, value)
                    if r == 0:
                        for run in cell_p.runs:
                            run.bold = True
            set_table_fixed(table, 9412)
            continue

        m = re.match(r"^(#{1,4})\s+(.*)$", stripped)
        if m:
            level = min(len(m.group(1)), 3)
            doc.add_heading(m.group(2), level=level)
            i += 1
            continue

        if stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, stripped[2:])
            i += 1
            continue

        if re.match(r"^\d+\.\s+", stripped):
            p = doc.add_paragraph(style="List Number")
            add_runs(p, re.sub(r"^\d+\.\s+", "", stripped))
            i += 1
            continue

        if stripped.startswith("> "):
            p = doc.add_paragraph(style="Quote")
            add_runs(p, stripped[2:])
            i += 1
            continue

        p = doc.add_paragraph()
        add_runs(p, stripped)
        i += 1

    if code_buf:
        flush_code()
    doc.save(DOCX_PATH)


def main() -> None:
    md_text = MD_PATH.read_text(encoding="utf-8")
    build_html(md_text)
    build_docx(md_text)
    print("built:", HTML_PATH.name, DOCX_PATH.name)


if __name__ == "__main__":
    sys.exit(main())
